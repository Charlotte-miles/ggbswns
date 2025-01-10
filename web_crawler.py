import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

# 目标URL
url = 'https://www.baidu.com'

try:
    # 发送HTTP请求
    proxies = {
        'http': None,
        'https': None
    }
    response = requests.get(url, proxies=proxies)
    response.encoding = 'utf-8'
    
    # 解析HTML
    soup = BeautifulSoup(response.text, 'html.parser')
    
    # 提取页面标题
    title = soup.title.text.strip()
    
    # 提取页面主要内容
    content = soup.find('div', id='wrapper').text.strip()
    
    # 创建Word文档
    doc = Document()
    doc.add_heading(title, 0)
    
    # 添加正文内容
    doc.add_paragraph(content)
    
    # 保存文档
    doc.save(f'{title}.docx')
    print(f'文档已成功保存为：{title}.docx')
    
except requests.exceptions.RequestException as e:
    print(f'网络请求出错：{e}')
except AttributeError as e:
    print(f'解析HTML出错：{e}')
except Exception as e:
    print(f'发生未知错误：{e}')
